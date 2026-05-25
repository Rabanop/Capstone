import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("La librería 'python-docx' no está instalada.")
    print("Por favor, instálala usando: pip install python-docx")
    sys.exit(1)

def main():
    doc = Document()
    
    # Título principal
    title = doc.add_heading('Informe Educativo: Sistema Dinámico de Agendamiento (VBS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Este documento explica de forma sencilla y educativa cómo funciona el simulador de agendamiento portuario y cómo el modelo dinámico resuelve los problemas de congestión.")
    
    # Sección 1
    doc.add_heading('1. ¿Qué problema estamos resolviendo?', level=1)
    p = doc.add_paragraph()
    p.add_run('Imagina que el puerto es un restaurante muy popular. ').bold = True
    p.add_run('En el modelo ').italic = False
    p.add_run('Estático (Tradicional)').bold = True
    p.add_run(', los camiones llegan a la hora que quieren sin importar si el patio está lleno. Esto es equivalente a ir al restaurante sin reserva: si no hay mesas, todos tienen que hacer fila en la acera, creando caos y bloqueando el tráfico.')
    
    p2 = doc.add_paragraph()
    p2.add_run('El modelo ').italic = False
    p2.add_run('Dinámico (Inteligente)').bold = True
    p2.add_run(' funciona como un sistema de reservas en tiempo real. Si el patio está a punto de llenarse (más del 60%) o la fila en la puerta es muy larga, el sistema contacta al camión antes de que salga de su origen y le pide que espere. De esta forma, la espera se realiza cómodamente en el origen y no haciendo fila en la puerta del puerto.')

    # Sección 2
    doc.add_heading('2. ¿Cómo solucionamos el problema técnico (El "Bug" de Congestión)?', level=1)
    doc.add_paragraph('Al principio, la simulación mostraba resultados idénticos para ambos modelos. Esto ocurrió porque el sistema evaluaba si había congestión al instante de iniciar el día (cuando todo estaba vacío), por lo que siempre dejaba pasar a todos.')
    doc.add_paragraph('Se corrigió implementando un control en tiempo real: ahora el sistema evalúa la congestión en el milisegundo exacto en que el camión está por salir. Esto demostró la efectividad del modelo "Pull" (Lean/Gemba).')
    
    # Sección 3
    doc.add_heading('3. Código Fuente Explicado', level=1)
    doc.add_paragraph('A continuación, incluimos el código fundamental de simulación que hace que esto sea posible:')
    
    code = """
def truck_process(self, truck_id: str, total_val_time: float, has_error: bool, scheduled_arrival: float):
    # El camión espera en su origen hasta la hora de su cita original
    if scheduled_arrival > self.env.now:
        yield self.env.timeout(scheduled_arrival - self.env.now)
        
    actual_arrival = scheduled_arrival
    
    # Escenario Dinámico: VBS actúa como sistema Pull (Lean)
    if self.is_dynamic:
        # Si la puerta tiene más del doble de su capacidad en fila, o el patio está al 60%
        # Retenemos al camión en su origen
        while len(self.gate.queue) > (self.gate.capacity * 2) or self.get_yard_occupancy() > 0.60:
            yield self.env.timeout(10.0) # Revisa de nuevo en 10 mins
            actual_arrival += 10.0
            
    # Llegada física al terminal
    arrival_time = self.env.now
    
    # Procesos de Puerta y Patio (Simulación física)
    # ...
    """
    code_para = doc.add_paragraph(code)
    code_para.style = 'Macro Text' # Simple monospaced style usually built in
    
    doc.add_heading('4. Conclusión Gemba', level=1)
    doc.add_paragraph('El control inteligente del inventario (camiones) antes de que ingresen al proceso elimina el "Work in Progress" innecesario, reduciendo horas de espera a minutos.')

    file_name = 'Informe_Educativo.docx'
    doc.save(file_name)
    print(f"El documento Word se ha generado exitosamente: {os.path.abspath(file_name)}")

if __name__ == "__main__":
    main()
